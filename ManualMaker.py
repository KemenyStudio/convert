import streamlit as st
import os
import sys
import time
import ffmpeg
from dotenv import load_dotenv
from azure.cognitiveservices.speech import SpeechConfig, AudioConfig, SpeechRecognizer, ResultReason
from azure.cognitiveservices.speech import CancellationReason
import azure.cognitiveservices.speech as speechsdk
import cv2
from docx import Document
from docx.shared import Inches
import tempfile
import shutil

load_dotenv()

API_KEY = os.getenv("API_KEY")
REGION = os.getenv("REGION")
LANGUAGE = os.getenv("LANGUAGE")

class AzureSpeechService:
    def __init__(self, speech_key, service_region):
        self.speech_key = speech_key
        self.service_region = service_region
        self.speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)

    def validate_speech_service(self):
        speech_recognizer = speechsdk.SpeechRecognizer(speech_config=self.speech_config)
        result = speech_recognizer.recognize_once_async().get()
        if result.reason == ResultReason.Canceled:
            cancellation_details = result.cancellation_details
            if cancellation_details.reason == CancellationReason.Error:
                if "401" in cancellation_details.error_details:
                    st.error("Invalid subscription key or region.")
                    return False
                else:
                    st.error(f"Error: {cancellation_details.error_details}")
                    return False
        st.success("Subscription key and region are valid.")
        return True

    def recognize(self, speech_recognizer: speechsdk.SpeechRecognizer) -> list:
        st.text("Starting speech recognition...")
        done = False
        all_results = list()

        def stop_cb(evt):
            speech_recognizer.stop_continuous_recognition()
            nonlocal done
            done = True

        def handle_final_result(evt):
            nonlocal all_results
            all_results.append(evt.result.text)

        speech_recognizer.recognized.connect(handle_final_result)
        speech_recognizer.session_stopped.connect(stop_cb)
        speech_recognizer.canceled.connect(stop_cb)

        speech_recognizer.start_continuous_recognition()
        while not done:
            time.sleep(5)

        return all_results

    def transcribe_audio(self, audio_file, language=LANGUAGE):
        self.speech_config.speech_recognition_language = language
        self.speech_config.set_property(speechsdk.PropertyId.SpeechServiceConnection_InitialSilenceTimeoutMs, '600000')
        audio_config = AudioConfig(filename=audio_file)
        speech_recognizer = SpeechRecognizer(speech_config=self.speech_config, audio_config=audio_config)

        all_results = self.recognize(speech_recognizer)

        if all_results:
            return " ".join(all_results)
        else:
            return "No speech could be recognized."

def extract_audio(video_file):
    audio_file = os.path.splitext(video_file)[0] + "_audio.wav"
    stream = ffmpeg.input(video_file)
    stream = ffmpeg.output(stream, audio_file, acodec='pcm_s16le', ac=1, ar=16000)
    ffmpeg.run(stream)
    return audio_file

def extract_frames(video_file, interval=10):
    video = cv2.VideoCapture(video_file)
    frames = []
    count = 0
    while True:
        success, image = video.read()
        if not success:
            break
        if count % (30 * interval) == 0:  # Extract frame every 'interval' seconds
            frames.append(image)
        count += 1
    video.release()
    return frames

def save_frames(frames):
    frame_paths = []
    for i, frame in enumerate(frames):
        path = os.path.join(tempfile.gettempdir(), f'frame_{i}.jpg')
        cv2.imwrite(path, frame)
        frame_paths.append(path)
    return frame_paths

def create_manual(transcription, frame_paths):
    doc = Document()
    doc.add_heading('User Guide', 0)

    paragraphs = transcription.split('. ')
    for i, paragraph in enumerate(paragraphs):
        doc.add_paragraph(paragraph + '.')
        if i < len(frame_paths):
            doc.add_picture(frame_paths[i], width=Inches(6))

    output_path = os.path.join(tempfile.gettempdir(), 'user_guide.docx')
    doc.save(output_path)
    return output_path

def process_video(video_file):
    with st.spinner('Extracting audio...'):
        audio_file = extract_audio(video_file)

    with st.spinner('Transcribing audio...'):
        azure_speech_service = AzureSpeechService(speech_key=API_KEY, service_region=REGION)
        transcription = azure_speech_service.transcribe_audio(audio_file)

    with st.spinner('Extracting frames...'):
        frames = extract_frames(video_file)
        frame_paths = save_frames(frames)

    with st.spinner('Creating user guide...'):
        user_guide_path = create_manual(transcription, frame_paths)

    st.success("Processing complete. User guide has been generated.")
    st.text("Transcription:")
    st.write(transcription)

    return user_guide_path

def main():
    st.title('Large Video Processing and Manual Generation')

    video_path = st.text_input("Enter the full path to your video file:")
    
    if video_path and os.path.exists(video_path):
        file_details = {"FileName": os.path.basename(video_path), "FileSize": os.path.getsize(video_path)}
        st.write(file_details)
        
        if st.button("Process Video"):
            # Process the video
            user_guide_path = process_video(video_path)
            
            # Offer the user guide for download
            with open(user_guide_path, "rb") as file:
                btn = st.download_button(
                    label="Download User Guide",
                    data=file,
                    file_name="user_guide.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            # Clean up temporary files
            os.unlink(user_guide_path)
            for file in os.listdir(tempfile.gettempdir()):
                if file.startswith('frame_') and file.endswith('.jpg'):
                    os.unlink(os.path.join(tempfile.gettempdir(), file))
    elif video_path:
        st.error("The specified file does not exist. Please check the path and try again.")

if __name__ == "__main__":
    main()